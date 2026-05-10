from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'docs' / 'DeployFlow_Final_Report_Draft.docx'


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Code']
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(8.5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(90, 90, 90)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles['Normal'].font.name = 'Arial'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    styles['Normal'].font.size = Pt(10)

    for name, size, color in [('Title', 24, RGBColor(32, 39, 56)), ('Heading 1', 16, RGBColor(39, 53, 86)), ('Heading 2', 13, RGBColor(39, 53, 86))]:
        style = styles[name]
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    if 'Code' not in [s.name for s in styles]:
        code_style = styles.add_style('Code', 1)
        code_style.font.name = 'Consolas'
        code_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        code_style.font.size = Pt(8.5)


def add_mapping_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Real deployment concept', 'Graph coloring concept']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade_cell(table.rows[0].cells[i], 'D9E2F3')
    rows = [
        ('Service deployment', 'Vertex'),
        ('Unsafe parallel relation', 'Edge'),
        ('Deployment window', 'Color'),
        ('Safe release plan', 'Valid graph coloring'),
    ]
    for a, b in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], a)
        set_cell_text(cells[1], b)


def add_contribution_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Team member', 'Contribution']):
        set_cell_text(table.rows[0].cells[i], h, True)
        shade_cell(table.rows[0].cells[i], 'E2F0D9')
    rows = [
        ('Artem Pasichnyk', 'Algorithm implementation, planner rules, testing, report sections.'),
        ('Yaroslav Kondratenko', 'Web interface, product design, screenshots, presentation sections.'),
    ]
    for a, b in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], a)
        set_cell_text(cells[1], b)


def build():
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DeployFlow\nSoftware Deployment Window Planner Using Graph Coloring')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(32, 39, 56)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Prepared by: Artem Pasichnyk and Yaroslav Kondratenko\n').bold = True
    p.add_run('Topic: Map Coloring Problem / Graph Coloring\n')
    p.add_run('Repository: replace with clickable GitHub repository link before submission')

    doc.add_page_break()

    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'DeployFlow is a software deployment window planner that applies graph coloring to daily release coordination. '
        'Each planned service deployment is modeled as a vertex. A conflict edge is created when two deployments should not run in parallel because they share an owner, use the same infrastructure, have a dependency relationship, or both have high operational risk. Deployment windows are modeled as colors. A valid coloring becomes a safe release plan in which conflicting deployments are placed into different time windows. The project implements the classical graph coloring backtracking algorithm in Java and extends it with practical improvements such as most-constrained-vertex ordering, degree-based tie breaking, least-constraining color choice, forward checking, and chronological dependency validation.'
    )
    doc.add_paragraph('Keywords: graph coloring, map coloring, backtracking, deployment planning, release scheduling, constraint satisfaction')

    doc.add_heading('1. Introduction and Problem Description', level=1)
    doc.add_paragraph(
        'The classical map coloring problem asks whether regions of a map can be colored so that adjacent regions do not share the same color. Algorithmically, the map is transformed into a graph: regions become vertices, shared borders become edges, and colors become labels assigned to vertices.'
    )
    doc.add_paragraph(
        'DeployFlow uses the same idea for software deployment planning. Engineering teams often have several service updates planned for one day. Running all of them at the same time is unsafe when releases affect the same infrastructure, are owned by the same team, or depend on each other. DeployFlow converts the daily deployment list into a conflict graph and assigns each deployment to a release window.'
    )
    add_mapping_table(doc)

    doc.add_heading('2. Classical Algorithm for the Problem', level=1)
    doc.add_paragraph(
        'The classical algorithm for graph coloring is recursive backtracking. It processes vertices one by one and tries to assign one of the available colors. A color is safe if none of the already colored adjacent vertices has the same color. If a vertex cannot receive any available color, the algorithm backtracks to a previous vertex and tries another color.'
    )
    doc.add_paragraph('The algorithm is complete: if a valid coloring exists for the chosen number of colors, backtracking will eventually find it. Its worst-case search space is exponential because it may need to explore many color combinations.')

    doc.add_heading('3. Pseudocode of the Classical Algorithm', level=1)
    add_code(doc, '''FUNCTION COLOR_VERTEX(vertexIndex):\n    IF vertexIndex = numberOfVertices:\n        RETURN true\n\n    FOR color FROM 1 TO numberOfColors:\n        IF IS_SAFE(vertexIndex, color):\n            colors[vertexIndex] = color\n\n            IF COLOR_VERTEX(vertexIndex + 1):\n                RETURN true\n\n            colors[vertexIndex] = 0\n\n    RETURN false\n\nFUNCTION IS_SAFE(vertexIndex, color):\n    FOR each neighbor of vertexIndex:\n        IF colors[neighbor] = color:\n            RETURN false\n    RETURN true''')

    doc.add_heading('4. Java Implementation of the Algorithm', level=1)
    doc.add_paragraph('The algorithm is implemented in src/com/deployflow/core/algorithm/GraphColoringSolver.java.')
    add_code(doc, '''private boolean classicColorVertex(int vertex, boolean[][] graph, int[] colors,\n                                   int colorCount, PlannerMetrics metrics,\n                                   List<String> trace) {\n    metrics.countRecursiveCall();\n    if (vertex == colors.length) {\n        return true;\n    }\n    for (int color = 1; color <= colorCount; color++) {\n        if (isClassicSafe(vertex, color, graph, colors, metrics)) {\n            colors[vertex] = color;\n            if (classicColorVertex(vertex + 1, graph, colors, colorCount, metrics, trace)) {\n                return true;\n            }\n            colors[vertex] = 0;\n            metrics.countBacktrack();\n        }\n    }\n    return false;\n}''')
    add_code(doc, '''private boolean isClassicSafe(int vertex, int color, boolean[][] graph,\n                              int[] colors, PlannerMetrics metrics) {\n    metrics.countSafetyCheck();\n    for (int other = 0; other < graph.length; other++) {\n        if (graph[vertex][other] && colors[other] == color) {\n            return false;\n        }\n    }\n    return true;\n}''')

    doc.add_heading('5. Real-World Situation', level=1)
    doc.add_paragraph('The selected real-world situation is software deployment planning. A development team may need to deploy multiple services on the same day. Some deployments cannot safely run in the same window.')
    add_bullets(doc, [
        'Checkout Web depends on Payments Gateway and Inventory Service.',
        'Payments Gateway and Identity API share critical identity/payment infrastructure.',
        'Two high-risk deployments should not run at the same time.',
        'Services owned by the same team should not be deployed in parallel because the same people monitor them.',
    ])

    doc.add_heading('6. Main Program Implementation', level=1)
    doc.add_paragraph('The driver program is implemented in src/com/deployflow/web/DeployFlowApp.java. It starts a Java HTTP server, serves the web interface, exposes API endpoints, and calls the graph coloring planner.')
    add_code(doc, '''public static void main(String[] args) throws Exception {\n    DeployFlowApp app = new DeployFlowApp();\n    if (args.length > 0 && "--demo".equalsIgnoreCase(args[0])) {\n        app.runConsoleDemo();\n        return;\n    }\n    int port = args.length > 0 ? parsePort(args[0]) : DEFAULT_PORT;\n    app.startServer(port);\n}''')
    add_code(doc, '''private void handlePlan(HttpExchange exchange) throws IOException {\n    String body = new String(exchange.getRequestBody().readAllBytes(), StandardCharsets.UTF_8);\n    Map<String, Object> request = Json.asObject(Json.parse(body));\n    List<DeploymentTask> tasks = parseTasks(request);\n    PlannerOptions options = PlannerOptions.fromMap(request);\n    Map<String, Object> result = planner.plan(tasks, options);\n    sendJson(exchange, 200, result);\n}''')

    doc.add_heading('7. Improved Algorithm', level=1)
    doc.add_paragraph('The improved algorithm keeps the same graph coloring foundation but reduces unnecessary search and adds deployment-specific chronological constraints.')
    add_numbered(doc, [
        'Most constrained vertex first: choose the uncolored deployment with the fewest currently available windows.',
        'Degree tie-breaker: choose the deployment with the most conflict edges.',
        'Risk tie-breaker: choose the higher-risk deployment first.',
        'Least-constraining color: prefer the window that leaves more choices for neighboring deployments.',
        'Forward checking: after assigning a deployment, check whether every uncolored deployment still has a possible window.',
        'Dependency order validation: if deployment A depends on deployment B, B must be assigned to an earlier window.',
    ])
    add_code(doc, '''FUNCTION IMPROVED_COLORING():\n    IF all vertices are colored:\n        RETURN true\n\n    vertex = SELECT_MOST_CONSTRAINED_VERTEX()\n    orderedColors = ORDER_COLORS_BY_LEAST_CONSTRAINING_VALUE(vertex)\n\n    FOR each color IN orderedColors:\n        IF IS_DEPLOYMENT_SAFE(vertex, color):\n            colors[vertex] = color\n\n            IF EVERY_UNCOLORED_VERTEX_HAS_A_VALID_COLOR():\n                IF IMPROVED_COLORING():\n                    RETURN true\n\n            colors[vertex] = 0\n\n    RETURN false''')

    doc.add_heading('8. Screenshots of Program Execution', level=1)
    dashboard = ROOT / 'docs' / 'screenshots' / '01-dashboard.png'
    plan = ROOT / 'docs' / 'screenshots' / '02-release-plan-report.png'
    doc.add_heading('Part 2.1. Screenshots: Dashboard', level=2)
    doc.add_picture(str(dashboard), width=Inches(6.5))
    add_caption(doc, 'Figure 1. DeployFlow dashboard with selected deployment queue and planner controls.')
    doc.add_heading('Part 2.2. Screenshots: Generated Release Plan', level=2)
    doc.add_picture(str(plan), width=Inches(5.5))
    add_caption(doc, 'Figure 2. Generated release timeline, conflict graph, runbook, and planner metrics.')
    doc.add_heading('Part 2.3. Console Execution', level=2)
    add_code(doc, '''DeployFlow console demo\nSolved: true\nMessage: Deployment plan generated.\nAlgorithm: Improved MRV + degree backtracking\nWindows used: 4\n\nWindow 1 09:30-10:15\n  - Feature Flags [Platform, Low]\n\nWindow 2 10:15-11:00\n  - Identity API [Platform, High]\n  - Inventory Service [Commerce, Medium]\n\nWindow 3 11:00-11:45\n  - Payments Gateway [Payments, Critical]\n  - Orders API [Commerce, Medium]\n  - Search Indexer [Data, Medium]\n\nWindow 4 11:45-12:30\n  - Checkout Web [Commerce, High]\n  - Notifications Worker [Messaging, Low]''')

    doc.add_heading('9. Contribution of Each Team Member', level=1)
    doc.add_paragraph('Replace this section with the final confirmed effort analysis before submission.')
    add_contribution_table(doc)

    doc.add_heading('10. One-Page Thesis / Abstract Description', level=1)
    doc.add_paragraph(
        'Software deployments are often planned under time pressure. When several service updates are scheduled for the same day, teams must decide which deployments can run in parallel and which must be separated. Manual planning becomes unreliable when services share infrastructure, depend on each other, or require the same engineering team for monitoring.'
    )
    doc.add_paragraph(
        'This project presents DeployFlow, a software deployment window planner based on graph coloring. Each deployment is represented as a graph vertex. A conflict edge is added between two vertices when the corresponding deployments cannot safely run at the same time. Reasons for conflict include shared team ownership, shared infrastructure, dependency relations, and high operational risk. Each deployment window is represented as a color. Therefore, producing a safe deployment plan becomes equivalent to finding a valid graph coloring where connected vertices receive different colors.'
    )
    doc.add_paragraph(
        'The project implements the classical recursive backtracking algorithm for graph coloring in Java. It then improves the classical method using most-constrained-vertex ordering, degree-based tie breaking, least-constraining color selection, forward checking, and chronological dependency validation. The improved algorithm is used in a working product interface that allows users to select planned deployments, configure available windows, generate a release timeline, inspect the conflict graph, and follow an operator runbook.'
    )

    doc.add_heading('11. Project Evaluation and Final Thoughts', level=1)
    doc.add_paragraph(
        'DeployFlow successfully applies the graph coloring model to a real-world scheduling problem. The product converts a deployment list into a concrete release timeline and exposes the reasoning behind the plan. The conflict graph and algorithm metrics make the result transparent: users can see why services were separated and how much search the solver performed.'
    )
    doc.add_paragraph(
        'The main limitation is that real deployment planning may include additional constraints such as stakeholder availability, external maintenance windows, rollback risk, and environment-specific rules. These can be added as additional graph edges or domain constraints.'
    )

    doc.add_heading('References', level=1)
    refs = [
        'T. R. Jensen and B. Toft, Graph Coloring Problems. New York: Wiley-Interscience, 1995.',
        'R. Diestel, Graph Theory. Berlin: Springer.',
        'A. Levitin, Introduction to the Design and Analysis of Algorithms. Boston: Pearson.',
        'Oracle, Java SE 21 API Specification: jdk.httpserver module.',
        'MDN Web Docs, HTML, CSS, JavaScript, and Fetch API documentation.',
    ]
    for ref in refs:
        doc.add_paragraph(ref, style='List Number')

    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    build()
